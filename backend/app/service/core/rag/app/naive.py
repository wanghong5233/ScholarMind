#
#  Copyright 2025 The InfiniFlow Authors. All Rights Reserved.
#
#  Licensed under the Apache License, Version 2.0 (the "License");
#  you may not use this file except in compliance with the License.
#  You may obtain a copy of the License at
#
#      http://www.apache.org/licenses/LICENSE-2.0
#
#  Unless required by applicable law or agreed to in writing, software
#  distributed under the License is distributed on an "AS IS" BASIS,
#  WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
#  See the License for the specific language governing permissions and
#  limitations under the License.
#

import logging
from tika import parser
from io import BytesIO
from docx import Document
from timeit import default_timer as timer
import re
from service.core.deepdoc.parser.pdf_parser import PlainParser
from service.core.rag.nlp import rag_tokenizer, naive_merge, tokenize_table, tokenize_chunks, find_codec, concat_img, \
    naive_merge_docx, tokenize_chunks_docx
from service.core.deepdoc.parser import PdfParser, ExcelParser, DocxParser, HtmlParser, JsonParser, MarkdownParser, TxtParser
from service.core.rag.utils import num_tokens_from_string
from PIL import Image
from functools import reduce
from markdown import markdown
from docx.image.exceptions import UnrecognizedImageError, UnexpectedEndOfFileError, InvalidImageStreamError


class Docx(DocxParser):
    def __init__(self):
        pass

    def get_picture(self, document, paragraph):
        img = paragraph._element.xpath('.//pic:pic')
        if not img:
            return None
        img = img[0]
        embed = img.xpath('.//a:blip/@r:embed')[0]
        related_part = document.part.related_parts[embed]
        try:
            image_blob = related_part.image.blob
        except UnrecognizedImageError:
            logging.info("Unrecognized image format. Skipping image.")
            return None
        except UnexpectedEndOfFileError:
            logging.info("EOF was unexpectedly encountered while reading an image stream. Skipping image.")
            return None
        except InvalidImageStreamError:
            logging.info("The recognized image stream appears to be corrupted. Skipping image.")
            return None
        try:
            image = Image.open(BytesIO(image_blob)).convert('RGB')
            return image
        except Exception:
            return None

    def __clean(self, line):
        line = re.sub(r"\u3000", " ", line).strip()
        return line

    def __call__(self, filename, binary=None, from_page=0, to_page=100000):
        """
        实现了Docx文档的解析流程。
        该方法作为Docx解析器的主要入口点，按顺序执行以下步骤：
        1. 使用python-docx库加载文档。
        2. 遍历文档中的所有段落（paragraphs）。
        3. 对每个段落进行清洗，提取文本内容，并尝试关联段落中的图片。
           - 特别处理了标题样式（Caption），将其与前一张图片关联。
           - 将文本和关联的图片作为一个元组存入列表。
        4. 通过检查XML中的'lastRenderedPageBreak'或'w:br'标签来模拟分页计数。
        5. 遍历文档中的所有表格（tables）。
        6. 将每个表格转换为HTML字符串格式。
        7. 返回处理好的文本-图片行列表和HTML表格列表。
        """
        self.doc = Document(
            filename) if not binary else Document(BytesIO(binary))
        pn = 0
        lines = []
        last_image = None
        #分页处理，将docx文件分页
        for p in self.doc.paragraphs:
            if pn > to_page:
                break
            if from_page <= pn < to_page:
                if p.text.strip():
                    if p.style and p.style.name == 'Caption':
                        former_image = None
                        if lines and lines[-1][1] and lines[-1][2] != 'Caption':
                            former_image = lines[-1][1].pop()
                        elif last_image:
                            former_image = last_image
                            last_image = None
                        lines.append((self.__clean(p.text), [former_image], p.style.name))
                    else:
                        current_image = self.get_picture(self.doc, p)
                        image_list = [current_image]
                        if last_image:
                            image_list.insert(0, last_image)
                            last_image = None
                        lines.append((self.__clean(p.text), image_list, p.style.name if p.style else ""))
                else:
                    if current_image := self.get_picture(self.doc, p):
                        if lines:
                            lines[-1][1].append(current_image)
                        else:
                            last_image = current_image
            for run in p.runs:
                if 'lastRenderedPageBreak' in run._element.xml:
                    pn += 1
                    continue
                if 'w:br' in run._element.xml and 'type="page"' in run._element.xml:
                    pn += 1
        new_line = [(line[0], reduce(concat_img, line[1]) if line[1] else None) for line in lines]

        tbls = []
        for tb in self.doc.tables:
            html = "<table>"
            for r in tb.rows:
                html += "<tr>"
                i = 0
                while i < len(r.cells):
                    span = 1
                    c = r.cells[i]
                    for j in range(i + 1, len(r.cells)):
                        if c.text == r.cells[j].text:
                            span += 1
                            i = j
                    i += 1
                    html += f"<td>{c.text}</td>" if span == 1 else f"<td colspan='{span}'>{c.text}</td>"
                html += "</tr>"
            html += "</table>"
            tbls.append(((None, html), ""))
        return new_line, tbls


class Pdf(PdfParser):
    def __call__(self, filename, binary=None, from_page=0,
                 to_page=100000, zoomin=3, callback=None):
        """
        实现了PDF文档的完整解析流程（Pipeline）。

        该方法作为PDF解析器的“总调度师”，按顺序执行父类`PdfParser`中定义的各个原子操作，
        将原始PDF文件转化为结构化的文本块和表格数据。

        处理流程:
        1. `__images__`: (OCR步骤) 将PDF页面渲染成图片，并利用OCR引擎提取所有文字块及其位置信息。
        2. `_layouts_rec`: (版面分析) 识别每个文字块的逻辑功能，如标题、段落、列表等。
        3. `_table_transformer_job`: (表格分析) 使用Transformer模型深入分析表格结构。
        4. `_text_merge`: (初步文本合并) 智能地将属于同一逻辑段落的文本块合并在一起。
        5. `_extract_table_figure`: (提取图表) 将分析出的表格和图片作为独立元素提取出来。
        6. `_concat_downward`: (最终合并) 执行更精细的向下合并策略，形成最终的文本段落。

        Returns:
            tuple[list[tuple[str, str]], list[dict]]: 
            一个元组(tuple)，包含两个核心元素：
            
            - 第一个元素 (文本块列表): `list[tuple[str, str]]`
              这是一个列表，其中每个元素都是一个元组，代表一个从文档中解析出的逻辑文本块（如一个段落、一个标题）。
              每个元组的内部结构是 `(text, metadata_tag)`:
                - `text` (str): 文本块的纯文本内容。
                - `metadata_tag` (str): 一个描述该文本块属性的标签。
                  这个标签通常包含了版面分析的结果（如 'title', 'paragraph', 'list_item'）、
                  所在的页码等信息。这些元数据对于后续的RAG分块(chunking)和构建索引至关重要。
                  例如，可以根据这些标签决定不同的分块策略。

            - 第二个元素 (表格列表): `list[dict]`
              这是一个列表，其中每个元素都是一个字典，代表一个从文档中完整提取出的表格。
              每个字典通常包含以下键值对：
                - `'title'` (str): 表格的标题或说明。
                - `'html'` (dtr): 表格内容的HTML字符串表示形式。将表格转换为HTML格式
                  有助于保留其二维结构，便于后续处理或直接展示。
        """
        start = timer()
        first_start = start
        callback(msg="OCR started")
        self.__images__(
            filename if not binary else binary,
            zoomin,
            from_page,
            to_page,
            callback
        )
        callback(msg="OCR finished ({:.2f}s)".format(timer() - start))
        logging.info("OCR({}~{}): {:.2f}s".format(from_page, to_page, timer() - start))

        start = timer()
        self._layouts_rec(zoomin)
        callback(0.63, "Layout analysis ({:.2f}s)".format(timer() - start))

        start = timer()
        self._table_transformer_job(zoomin)
        callback(0.65, "Table analysis ({:.2f}s)".format(timer() - start))

        start = timer()
        self._text_merge()
        callback(0.67, "Text merged ({:.2f}s)".format(timer() - start))
        tbls = self._extract_table_figure(True, zoomin, True, True)
        # self._naive_vertical_merge()
        self._concat_downward()
        # self._filter_forpages()

        logging.info("layouts cost: {}s".format(timer() - first_start))
        return [(b["text"], self._line_tag(b, zoomin))
                for b in self.boxes], tbls


class Markdown(MarkdownParser):
    def __call__(self, filename, binary=None):
        """
        实现了Markdown(.md)文件的解析流程。
        该方法作为Markdown解析器的主要入口点，执行以下操作：
        1. 读取文件内容，并自动检测编码。
        2. `extract_tables_and_remainder`: (提取表格) 首先从文本中分离出所有Markdown格式的表格。
        3. (文本块处理) 对剩余的文本进行处理：
           - 如果一个部分（按行分割）过长，则将其对半切分。
           - 如果一行以'#'开头，则认为它是一个新的标题/章节。
           - 如果一行不是标题，则尝试将其与前一个标题合并，形成完整的章节内容。
        4. (表格转换) 将之前分离出的Markdown表格转换为HTML格式。
        5. 返回处理好的文本章节列表和HTML表格列表。
        """
        if binary:
            encoding = find_codec(binary)
            txt = binary.decode(encoding, errors="ignore")
        else:
            with open(filename, "r") as f:
                txt = f.read()
        remainder, tables = self.extract_tables_and_remainder(f'{txt}\n')
        sections = []
        tbls = []
        for sec in remainder.split("\n"):
            if num_tokens_from_string(sec) > 3 * self.chunk_token_num:
                sections.append((sec[:int(len(sec) / 2)], ""))
                sections.append((sec[int(len(sec) / 2):], ""))
            else:
                if sec.strip().find("#") == 0:
                    sections.append((sec, ""))
                elif sections and sections[-1][0].strip().find("#") == 0:
                    sec_, _ = sections.pop(-1)
                    sections.append((sec_ + "\n" + sec, ""))
                else:
                    sections.append((sec, ""))

        for table in tables:
            tbls.append(((None, markdown(table, extensions=['markdown.extensions.tables'])), ""))
        return sections, tbls


def chunk(filename, binary=None, from_page=0, to_page=100000,
          lang="Chinese", callback=None, **kwargs):
    """
    一个核心的“调度-解析-分块”总函数。

    该函数是整个RAG数据处理流程的入口点，它根据文件类型动态选择最合适的解析器，
    然后将解析出的内容进行合并、切分，最终生成结构化的、可用于向量化的数据块(chunks)。

    处理流程:
    1.  **文件类型路由**:
        通过正则表达式检查文件名后缀（如 `.pdf`, `.docx`, `.txt` 等）。
    
    2.  **调用解析器 (Parse)**:
        根据文件类型，实例化并调用对应的解析器类（如 `Pdf`, `Docx`, `TxtParser`）。
        解析器负责将原始文件内容转化为一个初步的、由文本行和表格组成的列表 `sections`。
        例如，对于PDF，它会执行OCR、版面分析等复杂操作。
    
    3.  **合并 (Merge)**:
        将解析器返回的零散的 `sections` 列表传递给 `naive_merge` 或 `naive_merge_docx` 函数。
        这些合并函数会根据语义（如换行符、标点）和设定的 `chunk_token_num`（块目标Token数），
        将多个连续的文本行智能地合并成一个更大的、更有意义的文本块。
    
    4.  **最终分词与格式化 (Tokenize & Format)**:
        将合并后的大文本块列表传递给 `tokenize_chunks` 或 `tokenize_chunks_docx`。
        这些函数会为每个文本块进行精细化处理，包括：
        - 进行粗粒度和细粒度的分词。
        - 附加文档标题、文件名等元数据。
        - 构建成一个标准的字典结构，为后续生成向量和存入数据库做准备。

    Args:
        filename (str): 待处理的文件名。
        binary (bytes, optional): 文件的二进制内容，如果提供，则直接处理内存中的数据，否则从`filename`读取。
        from_page (int): 对于PDF等分页文档，指定开始处理的页码。
        to_page (int): 对于PDF等分页文档，指定结束处理的页码。
        lang (str): 文档的主要语言，影响分词策略。
        callback (callable, optional): 一个回调函数，用于在处理过程中报告进度。
        **kwargs: 其他配置参数，如 `parser_config`。

    Returns:
        list[dict]: 
        一个列表，其中每个元素都是一个处理完成的chunk（数据块）的字典。
        每个字典都包含了文本内容、分词结果、元数据等，可以直接用于后续的向量化和索引。
    """

    is_english = lang.lower() == "english"  # is_english(cks)
    parser_config = kwargs.get(
        "parser_config", {
            "chunk_token_num": 128, "delimiter": "\n!?。；！？", "layout_recognize": "DeepDOC"})
    doc = {
        "docnm_kwd": filename,
        "title_tks": rag_tokenizer.tokenize(re.sub(r"\.[a-zA-Z]+$", "", filename))
    }
    doc["title_sm_tks"] = rag_tokenizer.fine_grained_tokenize(doc["title_tks"])
    res = []
    pdf_parser = None
    if re.search(r"\.docx$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        sections, tables = Docx()(filename, binary)
        res = tokenize_table(tables, doc, is_english)  # just for table

        callback(0.8, "Finish parsing.")
        st = timer()

        chunks, images = naive_merge_docx(
            sections, int(parser_config.get(
                "chunk_token_num", 128)), parser_config.get(
                "delimiter", "\n!?。；！？"))

        if kwargs.get("section_only", False):
            return chunks

        res.extend(tokenize_chunks_docx(chunks, doc, is_english, images))
        logging.info("naive_merge({}): {}".format(filename, timer() - st))
        return res

    elif re.search(r"\.pdf$", filename, re.IGNORECASE):
        pdf_parser = Pdf()
        if parser_config.get("layout_recognize", "DeepDOC") == "Plain Text":
            pdf_parser = PlainParser()
        sections, tables = pdf_parser(filename if not binary else binary, from_page=from_page, to_page=to_page,
                                      callback=callback)
        res = tokenize_table(tables, doc, is_english)

    elif re.search(r"\.xlsx?$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        excel_parser = ExcelParser()
        
        # 如果没有 binary 数据，从文件路径读取
        if binary is None:
            with open(filename, 'rb') as f:
                binary = f.read()
        
        if parser_config.get("html4excel"):
            sections = [(_, "") for _ in excel_parser.html(binary, 12) if _]
        else:
            sections = [(_, "") for _ in excel_parser(binary) if _]

    elif re.search(r"\.(txt|py|js|java|c|cpp|h|php|go|ts|sh|cs|kt|sql)$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        sections = TxtParser()(filename, binary,
                               parser_config.get("chunk_token_num", 128),
                               parser_config.get("delimiter", "\n!?;。；！？"))
        callback(0.8, "Finish parsing.")

    elif re.search(r"\.(md|markdown)$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        sections, tables = Markdown(int(parser_config.get("chunk_token_num", 128)))(filename, binary)
        res = tokenize_table(tables, doc, is_english)
        callback(0.8, "Finish parsing.")

    elif re.search(r"\.(htm|html)$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        sections = HtmlParser()(filename, binary)
        sections = [(_, "") for _ in sections if _]
        callback(0.8, "Finish parsing.")

    elif re.search(r"\.json$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        chunk_token_num = int(parser_config.get("chunk_token_num", 128))
        sections = JsonParser(chunk_token_num)(binary)
        sections = [(_, "") for _ in sections if _]
        callback(0.8, "Finish parsing.")

    elif re.search(r"\.doc$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        binary = BytesIO(binary)
        doc_parsed = parser.from_buffer(binary)
        if doc_parsed.get('content', None) is not None:
            sections = doc_parsed['content'].split('\n')
            sections = [(_, "") for _ in sections if _]
            callback(0.8, "Finish parsing.")
        else:
            callback(0.8, f"tika.parser got empty content from {filename}.")
            logging.warning(f"tika.parser got empty content from {filename}.")
            return []

    else:
        raise NotImplementedError(
            "file type not supported yet(pdf, xlsx, doc, docx, txt supported)")

    st = timer()
    chunks = naive_merge(
        sections, int(parser_config.get(
            "chunk_token_num", 128)), parser_config.get(
            "delimiter", "\n!?。；！？"))
    if kwargs.get("section_only", False):
        return chunks

    res.extend(tokenize_chunks(chunks, doc, is_english, pdf_parser))
    logging.info("naive_merge({}): {}".format(filename, timer() - st))
    return res


if __name__ == "__main__":
    import sys


    def dummy(prog=None, msg=""):
        pass


    chunk(sys.argv[1], from_page=0, to_page=10, callback=dummy)
